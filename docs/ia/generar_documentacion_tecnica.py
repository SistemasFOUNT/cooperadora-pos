from datetime import date
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_separator(doc):
    doc.add_paragraph('-' * 60)


def build_document():
    doc = Document()

    # Portada
    add_title(doc, 'DOCUMENTACION TECNICA DEL SISTEMA')
    add_title(doc, 'POS COOPERADORA - FACULTAD DE ODONTOLOGIA')
    add_subtitle(doc, 'Version: 1.0')
    add_subtitle(doc, f'Fecha: {date.today().strftime("%d/%m/%Y")}')
    add_subtitle(doc, 'Base: Protocolo skill-driven + spec-driven (v2.0)')
    doc.add_page_break()

    add_h1(doc, '1. Objetivo y alcance')
    doc.add_paragraph(
        'Este documento consolida la arquitectura tecnica, componentes, flujos principales, '
        'reglas de seguridad, esquema de datos y practicas operativas del sistema Cooperadora. '
        'El alcance incluye los modulos BOX, Postgrado, Odonto, Administracion, Contabilidad y '
        'el subsistema de financiamiento interno.'
    )

    add_h1(doc, '2. Resumen del sistema')
    add_bullets(doc, [
        'Framework principal: Laravel 11 sobre PHP 8.2+ (plataforma fijada a PHP 8.4.18 en composer).',
        'Base de datos principal: PostgreSQL.',
        'Front-end: Blade + Bootstrap + jQuery, assets con Vite/Tailwind.',
        'Auditoria: owen-it/laravel-auditing.',
        'Autorizacion por roles y permisos: spatie/laravel-permission.',
        'Panel y layout administrativo: jeroennoten/laravel-adminlte.',
        'Objetivo funcional: POS multi punto de venta con aislamiento estricto por dominio.'
    ])

    add_h1(doc, '3. Arquitectura de alto nivel')
    add_h2(doc, '3.1 Estilo arquitectonico')
    doc.add_paragraph(
        'Aplicacion monolitica modular. El enrutamiento web y API concentra la orquestacion, '
        'los controladores resuelven casos de uso por modulo y los servicios encapsulan logica '
        'de dominio especifica (contabilidad, facturacion, postgrado, odonto, box).'
    )

    add_h2(doc, '3.2 Capas tecnicas')
    add_bullets(doc, [
        'Capa de presentacion: vistas Blade, componentes, assets compilados con Vite.',
        'Capa de aplicacion: controladores en app/Http/Controllers y middlewares en app/Http/Middleware.',
        'Capa de dominio: modelos Eloquent y servicios en app/Services.',
        'Capa de persistencia: PostgreSQL, migraciones en database/migrations.',
        'Capa operativa: scripts de deployment y stack Docker para produccion/local estandarizada.'
    ])

    add_h2(doc, '3.3 Modulos funcionales principales')
    add_bullets(doc, [
        'BOX Cooperadora: cobros de productos, cuotas, bonos, odontologia y otros conceptos; inventario; pagos a proveedores; reportes.',
        'Postgrado: carreras, estudiantes, inscripciones, cobros por programa, reportes y certificados.',
        'Odonto: pacientes, agenda, tratamientos, POS y reportes.',
        'Admin: supervision consolidada, ingresos/egresos, libro caja, auditoria y autorizaciones.',
        'Contable: libros diario/caja/banco, plan de cuentas, estados y reportes.',
        'Financiamiento interno: clientes deudores, financiamientos, cuotas y documentos legales.'
    ])

    add_h1(doc, '4. Enrutamiento y contratos funcionales')
    add_h2(doc, '4.1 Rutas nucleares')
    add_bullets(doc, [
        'Dashboard dinamico por rol: admin -> admin.dashboard; usuario_box -> box.dashboard; usuario_postgrado -> postgrado.dashboard; usuario_odonto -> odonto.dashboard.',
        'Prefijo /box con middleware punto_venta + box_menu.',
        'Prefijo /postgrado con middleware punto_venta + postgrado_menu.',
        'Prefijo /odonto con middleware punto_venta + odonto_menu.',
        'Prefijo /admin con middleware admin + admin_menu.',
        'Prefijo /admin/contable que incluye routes/contable.php.'
    ])

    add_h2(doc, '4.2 API internas y endpoints especificos')
    add_bullets(doc, [
        'API autenticada para clientes deudores y financiamientos (/api/*).',
        'Rutas de documentos legales para compromiso de pago, registro de impresion e integridad.',
        'Rutas CRUD de clientes deudores y panel administrativo de financiamientos.'
    ])

    add_h1(doc, '5. Seguridad y aislamiento por punto de venta')
    add_h2(doc, '5.1 Middlewares registrados (bootstrap/app.php)')
    add_bullets(doc, [
        'role, permission, role_or_permission (Spatie).',
        'punto_venta (validacion de acceso por dominio).',
        'admin, admin_menu, box_menu, postgrado_menu, odonto_menu.'
    ])

    add_h2(doc, '5.2 Regla de aislamiento')
    doc.add_paragraph(
        'El middleware punto_venta permite acceso total a admin y restringe a cada usuario '
        'a su punto_venta_id. Controladores de BOX, Postgrado y Odonto refuerzan la validacion '
        'en consultas y acciones criticas con filtros por punto_venta_id.'
    )

    add_h2(doc, '5.3 Modelo de roles')
    add_bullets(doc, [
        'admin',
        'usuario_box',
        'usuario_postgrado',
        'usuario_odonto'
    ])

    add_h1(doc, '6. Modelo de datos (resumen tecnico)')
    add_h2(doc, '6.1 Tablas base de plataforma')
    add_bullets(doc, [
        'users, password_reset_tokens, sessions',
        'jobs, job_batches, failed_jobs',
        'cache, cache_locks'
    ])

    add_h2(doc, '6.2 Tablas de negocio principales')
    add_bullets(doc, [
        'branches, products, payment_methods, employees',
        'students y estudiantes (estructura historica/legacy)',
        'sales y sale_items (estructura base de migraciones)',
        'cash_movements, stock_movements',
        'puntos_venta, career_fee_config, cuotas_estudiantiles',
        'facturas, proveedores, pagos_proveedores, arqueos_caja',
        'cuentas_contables, asientos_contables, movimientos_contables',
        'clientes_deudores, financiamientos_odontologia, cuotas_financiamiento, documentos_legales',
        'configuracion_organizacion'
    ])

    add_h2(doc, '6.3 Consideracion tecnica de nomenclatura (legacy)')
    doc.add_paragraph(
        'Se observa coexistencia de nomenclatura en ingles/espanol. Ejemplo: migraciones base '
        'crean sales/sale_items, mientras modelos y consultas operativas usan ventas/items_venta. '
        'Esto requiere validacion explicita de esquema al migrar o desplegar en entornos nuevos.'
    )

    add_h1(doc, '7. Protocolo de cobro unificado (regla critica)')
    add_paragraph = doc.add_paragraph
    add_paragraph(
        'El proyecto define estandar obligatorio para todo flujo de cobro, independientemente del modulo. '
        'La paridad funcional es requisito de calidad y anti-regresion.'
    )
    add_bullets(doc, [
        'Modal de pago comun para todos los modulos equivalentes.',
        'Metodos estandar: efectivo, tarjeta, transferencia, mixto.',
        'Comprobantes estandar: ticket, factura_local, factura_fiscal.',
        'Funciones JS obligatorias y consistencia visual/UX transversal.',
        'Regla de oro: si cambia en un modulo, se replica igual en los demas modulos equivalentes.'
    ])

    add_h1(doc, '8. Facturacion e integracion fiscal')
    add_bullets(doc, [
        'Configuracion central en config/facturacion.php.',
        'Soporte de facturacion local habilitado por defecto.',
        'Integracion ARCA/AFIP configurable (testing/production) con endpoints diferenciados.',
        'Tipos de comprobante soportados: A, B, C.',
        'Numeracion por punto de venta para facturas locales.'
    ])

    add_h1(doc, '9. Infraestructura y despliegue')
    add_h2(doc, '9.1 Entornos')
    add_bullets(doc, [
        'Desarrollo: Windows + PostgreSQL.',
        'Produccion: Ubuntu Linux + PostgreSQL.'
    ])

    add_h2(doc, '9.2 Stack Docker (docker-compose.yml)')
    add_bullets(doc, [
        'app: contenedor PHP/Laravel.',
        'nginx: proxy HTTP/HTTPS y publicacion de assets.',
        'db: PostgreSQL 14 con volumen persistente.',
        'redis: cache/colas para escenarios de produccion.'
    ])

    add_h2(doc, '9.3 Scripts operativos')
    add_bullets(doc, [
        'deployment/deploy.sh, quick-deploy.sh y guias de post-deploy/health-check.',
        'Scripts de correccion operativa: fix-error-500.sh, fix-assets.sh, fix-htaccess.sh.'
    ])

    add_h1(doc, '10. Observabilidad, auditoria y cumplimiento')
    add_bullets(doc, [
        'Trazabilidad de cambios mediante Laravel Auditing.',
        'Vistas y rutas de auditoria para administracion.',
        'Controles de autorizacion y validacion por middleware y rol.',
        'Modulos contables para seguimiento de ingresos, egresos y libros legales.'
    ])

    add_h1(doc, '11. Calidad, pruebas y criterios de cierre')
    add_h2(doc, '11.1 Suite de pruebas identificada')
    add_bullets(doc, [
        'Feature: BoxAccessTest, AuditoriaFeatureTest, PuntosVentaReportesFeatureTest, Auth y Profile.',
        'Unit: ExampleTest base.',
        'Framework de pruebas: PHPUnit 11.'
    ])

    add_h2(doc, '11.2 Criterio operativo recomendado (segun protocolo)')
    add_numbered(doc, [
        'Reconocimiento de modulo y dependencias (SKILL-01).',
        'Validacion de aislamiento entre puntos de venta (SKILL-02).',
        'Cambio minimo y seguro (SKILL-03).',
        'Validacion de cobro unificado cuando aplique (SKILL-04).',
        'Pruebas de no-regresion y evidencia de cierre (SKILL-05/SKILL-06).'
    ])

    add_h1(doc, '12. Riesgos tecnicos y recomendaciones')
    add_bullets(doc, [
        'Riesgo de drift de esquema por coexistencia de nombres ingles/espanol en tablas.',
        'Riesgo de regression si cambios de cobro no se replican entre modulos equivalentes.',
        'Riesgo operativo si no se valida punto_venta_id en nuevas consultas/controladores.'
    ])
    add_numbered(doc, [
        'Definir y ejecutar plan de normalizacion de nomenclatura de tablas en entorno controlado.',
        'Agregar pruebas de integracion para flujos de cobro unificado por modulo.',
        'Incorporar checklist de aislamiento como gate obligatorio previo a merge.'
    ])

    add_h1(doc, '13. Inventario de archivos clave consultados')
    add_bullets(doc, [
        'docs/ia/PROTOCOLO_ANALISIS_SISTEMA_COOPERADORA.md',
        'ESTANDARES-PROYECTO.md',
        'routes/web.php y routes/contable.php',
        'bootstrap/app.php',
        'app/Http/Middleware/PuntoVentaMiddleware.php',
        'app/Models/User.php, Sale.php, SaleItem.php, Student.php',
        'database/migrations/*.php (tablas de negocio y plataforma)',
        'config/facturacion.php',
        'docker-compose.yml',
        'README.md y README_SISTEMA.md'
    ])

    add_separator(doc)
    doc.add_paragraph('Fin de documento tecnico.')

    output_path = 'docs/ia/DOCUMENTACION_TECNICA_SISTEMA_COOPERADORA.docx'
    doc.save(output_path)
    print(output_path)


if __name__ == '__main__':
    build_document()
